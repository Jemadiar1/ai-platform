"""
Servicio de generación de reportes profesionales.

Pipeline:
1. Renderizar gráficos como imágenes (matplotlib)
2. Renderizar HTML con Jinja2 + theme
3. Generar formatos: HTML, PDF (WeasyPrint), DOCX (python-docx), XLSX (openpyxl), CSV
4. Guardar en BD con metadata

Usar:
    from ai_platform.services.report_renderer import ReportRendererService
    renderer = ReportRendererService()
    result = renderer.render(tenant_id, report_spec, formats=["html", "pdf"])
"""

import base64
import csv
import io
import logging
import time
from dataclasses import asdict
from typing import Any

from docx.shared import Pt

from ai_platform.core.config import get_settings
from ai_platform.database import make_session
from ai_platform.models.db import GeneratedReport, UsageEvent
from ai_platform.services.report_models import (
    ChartType,
    ReportFormat,
    ReportSpec,
)

logger = logging.getLogger(__name__)
settings = get_settings()


class ReportRendererService:
    """
    Servicio de generación de reportes profesionales.
    """

    def render(
        self,
        tenant_id: str,
        report_spec: ReportSpec,
        formats: list[ReportFormat] | None = None,
    ) -> dict[str, bytes]:
        """
        Pipeline completo de renderizado.

        Args:
            tenant_id: Tenant propietario del reporte.
            report_spec: Especificación del reporte.
            formats: Formatos a generar (default: [HTML, PDF]).

        Returns:
            Dict {format: bytes} con los outputs generados.
        """
        start_time = time.time()

        if formats is None:
            formats = [ReportFormat.HTML, ReportFormat.PDF]

        # 1. Generar gráficos como imágenes
        chart_images = self._render_charts(report_spec)

        # 2. Renderizar HTML
        html_content = self._render_html(report_spec, chart_images)

        # 3. Generar formatos solicitados
        outputs: dict[str, bytes] = {"html": html_content.encode("utf-8")}

        if ReportFormat.PDF in formats:
            outputs["pdf"] = self._generate_pdf(html_content)

        if ReportFormat.DOCX in formats:
            outputs["docx"] = self._generate_docx(report_spec, chart_images)

        if ReportFormat.XLSX in formats:
            outputs["xlsx"] = self._generate_xlsx(report_spec)

        if ReportFormat.CSV in formats:
            outputs["csv"] = self._generate_csv(report_spec)

        # 4. Guardar en BD
        rendering_time = int((time.time() - start_time) * 1000)
        total_size = sum(len(v) for v in outputs.values())
        self._save_report(tenant_id, report_spec, formats, outputs, rendering_time, total_size)

        # 5. Log usage event
        self._log_usage(tenant_id, formats, rendering_time)

        logger.info(
            "report_rendered",
            tenant_id=tenant_id,
            title=report_spec.title,
            formats=[f.value for f in formats],
            charts=report_spec.chart_count,
            tables=report_spec.table_count,
            rendering_ms=rendering_time,
        )

        return outputs

    # =========================================================================
    # Chart rendering
    # =========================================================================

    def _render_charts(self, spec: ReportSpec) -> dict[str, bytes]:
        """Renderizar todos los gráficos como PNG en memoria."""
        import matplotlib

        matplotlib.use("Agg")  # Backend sin GUI
        import matplotlib.pyplot as plt

        outputs: dict[str, bytes] = {}

        for section in spec.sections:
            for chart in section.charts:
                fig, ax = plt.subplots(figsize=(chart.width / 100, chart.height / 100))

                try:
                    if chart.type == ChartType.BAR:
                        labels = [d.get("label", "") for d in chart.data]
                        values = [d.get("value", 0) for d in chart.data]
                        colors = chart.colors or ["#1a73e8", "#e84393", "#00b894", "#fdcb6e", "#6c5ce7"][: len(labels)]
                        ax.bar(labels, values, color=colors)

                    elif chart.type == ChartType.LINE:
                        labels = [d.get("label", "") for d in chart.data]
                        values = [d.get("value", 0) for d in chart.data]
                        color = chart.colors[0] if chart.colors else "#1a73e8"
                        ax.plot(labels, values, marker="o", color=color)
                        ax.fill_between(labels, values, alpha=0.1)

                    elif chart.type == ChartType.PIE:
                        labels = [d.get("label", "") for d in chart.data]
                        values = [d.get("value", 0) for d in chart.data]
                        colors = chart.colors or plt.cm.Set3.colors[: len(labels)]
                        ax.pie(values, labels=labels, colors=colors, autopct="%1.1f%%")

                    ax.set_title(chart.title)
                    ax.tick_params(axis="x", rotation=45)

                    buf = io.BytesIO()
                    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
                    buf.seek(0)
                    outputs[chart.id] = buf.getvalue()

                finally:
                    plt.close(fig)

        return outputs

    # =========================================================================
    # HTML rendering
    # =========================================================================

    def _render_html(self, spec: ReportSpec, chart_images: dict[str, bytes]) -> str:
        """Renderizar HTML con Jinja2 + theme CSS variables + gráficos embebidos."""
        from jinja2 import BaseLoader, Environment

        template = """<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <title>{{ title }}</title>
    <style>
        :root {
            --primary: {{ theme.primary_color }};
            --secondary: {{ theme.secondary_color }};
            --font: {{ theme.font_family }};
        }
        body { font-family: var(--font); margin: 0; padding: 20px; color: #333; }
        .header { background: var(--primary); color: white; padding: 20px; margin-bottom: 30px; border-radius: 4px; }
        .header h1 { margin: 0; }
        .section { margin: 30px 0; page-break-inside: avoid; }
        .section h2 { border-bottom: 2px solid var(--primary); padding-bottom: 8px; }
        .chart { text-align: center; margin: 20px 0; }
        .chart img { max-width: 100%; height: auto; }
        table { border-collapse: collapse; width: 100%; margin: 15px 0; }
        th { background: var(--primary); color: white; padding: 10px; text-align: left; }
        td { border: 1px solid #ddd; padding: 8px 10px; }
        tr:nth-child(even) { background: #f9f9f9; }
        .citation { font-size: 0.85em; color: #666; margin-top: 5px; }
        .footer { margin-top: 40px; padding-top: 20px; border-top: 1px solid #ddd; font-size: 0.8em; color: #999; }
    </style>
</head>
<body>
    <div class="header">
        <h1>{{ title }}</h1>
        <p>Audiencia: {{ audience }}</p>
        <p>Generado por {{ generated_by }} v{{ version }}</p>
    </div>

    {% for section in sections %}
    <div class="section">
        <h2>{{ section.title }}</h2>
        <div class="content">{{ section.content | safe }}</div>

        {% for chart in section.charts %}
        <div class="chart">
            <h3>{{ chart.title }}</h3>
            <img src="data:image/png;base64,{{ chart_b64[chart.id] }}"
                  alt="{{ chart.title }}" width="{{ chart.width }}">
        </div>
        {% endfor %}

        {% for table in section.tables %}
        <table>
            <thead><tr>{% for h in table.headers %}<th>{{ h }}</th>{% endfor %}</tr></thead>
            <tbody>
                {% for row in table.rows %}
                <tr>{% for cell in row %}<td>{{ cell }}</td>{% endfor %}</tr>
                {% endfor %}
            </tbody>
        </table>
        {% endfor %}

        {% for cite in section.citations %}
        <div class="citation">📎 {{ cite.text }} — {{ cite.source }}{% if cite.url %} ({{ cite.url }}){% endif %}</div>
        {% endfor %}
    </div>
    {% endfor %}

    <div class="footer">
        <p>{{ theme.company_name }} | {{ title }} | Generado {{ generated_by }}</p>
    </div>
</body>
</html>"""

        env = Environment(loader=BaseLoader())
        template_obj = env.from_string(template)

        # Convert chart images to base64 for embedding
        chart_b64 = {}
        for chart_id, img_bytes in chart_images.items():
            chart_b64[chart_id] = base64.b64encode(img_bytes).decode("utf-8")

        return template_obj.render(
            title=spec.title,
            audience=spec.audience,
            sections=spec.sections,
            theme=spec.theme,
            chart_b64=chart_b64,
            generated_by=spec.generated_by,
            version=spec.version,
        )

    # =========================================================================
    # PDF generation
    # =========================================================================

    def _generate_pdf(self, html_content: str) -> bytes:
        """Generar PDF desde HTML usando WeasyPrint."""
        try:
            from weasyprint import HTML

            return HTML(string=html_content).write_pdf()
        except ImportError:
            logger.warning("weasyprint not installed, skipping PDF generation")
            return b""

    # =========================================================================
    # DOCX generation
    # =========================================================================

    def _generate_docx(self, spec: ReportSpec, chart_images: dict[str, bytes]) -> bytes:
        """Generar DOCX desde ReportSpec usando python-docx."""
        try:
            from docx import Document
            from docx.shared import Inches

            doc = Document()

            style = doc.styles["Normal"]
            font = style.font
            font.name = spec.theme.font_family.split(",")[0].strip()
            font.size = Pt(11)

            doc.add_heading(spec.title, level=0)
            doc.add_paragraph(f"Audiencia: {spec.audience}")

            for section in spec.sections:
                doc.add_heading(section.title, level=1)
                doc.add_paragraph(section.content)

                for chart in section.charts:
                    if chart.id in chart_images:
                        doc.add_picture(
                            io.BytesIO(chart_images[chart.id]),
                            width=Inches(5),
                        )

                for table in section.tables:
                    t = doc.add_table(rows=1 + len(table.rows), cols=len(table.headers))
                    for i, header in enumerate(table.headers):
                        t.rows[0].cells[i].text = str(header)
                    for row_idx, row in enumerate(table.rows):
                        for col_idx, cell in enumerate(row):
                            t.rows[row_idx + 1].cells[col_idx].text = str(cell)

            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            return buf.getvalue()
        except ImportError:
            logger.warning("python-docx not installed, skipping DOCX generation")
            return b""

    # =========================================================================
    # Spreadsheet generation
    # =========================================================================

    def _generate_xlsx(self, spec: ReportSpec) -> bytes:
        """Generar XLSX desde ReportSpec usando openpyxl."""
        try:
            from openpyxl import Workbook

            wb = Workbook()

            for section in spec.sections:
                ws = wb.create_sheet(title=section.title[:31])
                ws.append([section.title])

                for table in section.tables:
                    ws.append(table.headers)
                    for row in table.rows:
                        ws.append(row)

                for chart in section.charts:
                    ws.append(["Gráfico: " + chart.title])
                    for d in chart.data:
                        ws.append([d.get("label", ""), d.get("value", "")])

            buf = io.BytesIO()
            wb.save(buf)
            buf.seek(0)
            return buf.getvalue()
        except ImportError:
            logger.warning("openpyxl not installed, skipping XLSX generation")
            return b""

    def _generate_csv(self, spec: ReportSpec) -> bytes:
        """Generar CSV desde ReportSpec."""

        buf = io.BytesIO()
        writer = csv.writer(buf)

        for section in spec.sections:
            writer.writerow([f"SECCIÓN: {section.title}"])
            writer.writerow([])

            for table in section.tables:
                writer.writerow(table.headers)
                for row in table.rows:
                    writer.writerow(row)
                writer.writerow([])

            for chart in section.charts:
                writer.writerow([f"GRÁFICO: {chart.title}"])
                for d in chart.data:
                    writer.writerow([d.get("label", ""), d.get("value", "")])
                writer.writerow([])

        buf.seek(0)
        return buf.getvalue()

    # =========================================================================
    # Persistence
    # =========================================================================

    def _save_report(
        self,
        tenant_id: str,
        spec: ReportSpec,
        formats: list[ReportFormat],
        outputs: dict[str, bytes],
        rendering_time_ms: int,
        total_size: int,
    ) -> None:
        """Guardar reporte en BD."""
        with make_session() as db:
            report = GeneratedReport(
                tenant_id=tenant_id,
                title=spec.title,
                audience=spec.audience,
                generated_formats=[f.value for f in formats],
                report_spec=asdict(spec),
                html_content=outputs.get("html", b"").decode("utf-8", errors="replace"),
                pdf_blob=outputs.get("pdf"),
                docx_blob=outputs.get("docx"),
                xlsx_blob=outputs.get("xlsx"),
                csv_content=outputs.get("csv", b"").decode("utf-8", errors="replace"),
                file_size_bytes=total_size,
                rendering_time_ms=rendering_time_ms,
            )
            db.add(report)
            db.commit()

    def _log_usage(self, tenant_id: str, formats: list[ReportFormat], rendering_time_ms: int) -> None:
        """Registrar usage event."""
        with make_session() as db:
            event = UsageEvent(
                tenant_id=tenant_id,
                module="report_renderer",
                event_type="report_render",
                tokens_used=0,
                cost_usd=0.0,
                extra_data={
                    "formats": [f.value for f in formats],
                    "rendering_ms": rendering_time_ms,
                },
            )
            db.add(event)
            db.commit()
